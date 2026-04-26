"""自检 V5 论文：统计字数、段落、图、表、参考文献条数。"""
from pathlib import Path
import re

from docx import Document

PATH = Path(r"c:\Users\ylq06\Desktop\毕业设计\论文\毕业设计论文 - V5.0_完整版.docx")
doc = Document(str(PATH))

para_count = len(doc.paragraphs)

# 字数统计：仅统计中英文/数字字符
total_text = ""
for p in doc.paragraphs:
    total_text += p.text
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                total_text += p.text

# 中英文字符
chinese = len(re.findall(r"[\u4e00-\u9fff]", total_text))
english_words = len(re.findall(r"[A-Za-z]+", total_text))
digits = len(re.findall(r"\d+", total_text))

# 字数 = 中文字数 + 英文单词数（按论文常用算法）
word_count = chinese + english_words

# 图：每个 inline_shape
img_count = sum(len(p.runs) for p in doc.paragraphs)
shapes = doc.inline_shapes
n_imgs = len(shapes)

n_tables = len(doc.tables)

# 章节统计：按"第X章" 出现次数
chapters = re.findall(r"第\d+章", total_text)
unique_chapters = set(chapters)

# 参考文献条数
refs = re.findall(r"^\[\d+\]", total_text, re.M) + re.findall(r"\[\d+\]\s+[A-Za-z\u4e00-\u9fff]", total_text)

# 找参考文献区域更精确
ref_section_started = False
ref_lines = []
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt == "参考文献":
        ref_section_started = True
        continue
    if ref_section_started:
        if txt.startswith("附录"):
            break
        if re.match(r"^\[\d+\]", txt):
            ref_lines.append(txt)

print("=" * 60)
print(f"V5 论文统计：{PATH.name}")
print("=" * 60)
print(f"段落数        : {para_count}")
print(f"表格数        : {n_tables}")
print(f"图片数        : {n_imgs}")
print(f"中文字符      : {chinese:,}")
print(f"英文单词      : {english_words:,}")
print(f"字数(中文+英文): {word_count:,}")
print(f"独立章节      : {sorted(unique_chapters)}")
print(f"参考文献条数  : {len(ref_lines)}")
print(f"文件大小      : {PATH.stat().st_size/1024:.1f} KB")
print("=" * 60)
print("\n《阳光学院本科生毕业论文（设计）撰写规范》要求：")
print(f"  - 软件设计类 ≥ 8000 字          : {'PASS' if word_count >= 8000 else 'FAIL'}  ({word_count:,})")
print(f"  - 参考文献 ≥ 10 篇              : {'PASS' if len(ref_lines) >= 10 else 'FAIL'}  ({len(ref_lines)})")
print(f"  - 图(技术流程图/架构图等) ≥ 5  : {'PASS' if n_imgs >= 5 else 'FAIL'}  ({n_imgs})")
